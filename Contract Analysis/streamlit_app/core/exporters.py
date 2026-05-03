"""Export helpers: redlined DOCX, JSON, Markdown."""
from __future__ import annotations

import io
import json
from typing import List

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from .state import Modification


def _add_strikethrough(run):
    rPr = run._element.get_or_add_rPr()
    rPr.append(OxmlElement("w:strike"))


def redlined_docx_bytes(contract_text: str, mods: List[Modification], title: str = "Redlined Contract") -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for para in contract_text.split("\n"):
        modified = False
        for mod in mods:
            if mod.original_text and mod.original_text.strip() and mod.original_text in para:
                modified = True
                p = doc.add_paragraph()
                run_orig = p.add_run(mod.original_text)
                _add_strikethrough(run_orig)
                run_new = p.add_run("\n" + mod.suggested_text)
                run_new.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                run_reason = p.add_run(f"\n[{mod.risk_level}] {mod.reason}")
                run_reason.italic = True
                run_reason.font.size = Pt(9)
                run_reason.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                break
        if not modified:
            doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def report_docx_bytes(report_md: str, title: str = "Contract Review Report") -> bytes:
    """Render the markdown report as a simple DOCX (paragraph per line)."""
    doc = Document()
    doc.add_heading(title, level=1)
    for line in report_md.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def report_json_bytes(report_json: dict) -> bytes:
    return json.dumps(report_json, indent=2, default=str).encode("utf-8")
